import hashlib
import os
from docx import Document


def calc_sha1_file(filepath):
    """Calculate the hash of the file"""
    return hashlib.sha1(filepath.encode()).hexdigest()


def listfiles(path):
    """List the files at a specific location"""
    list_of_files = list()
    for root, directory, files in os.walk(path):
        for filename in files:
            list_of_files.append(os.path.join(root, filename))
    return list_of_files


def generate_list_hash(location):
    """Generate the hashes for all files at this path"""
    files = listfiles(location)
    liststring = list()
    for filename in files:
        sha1 = calc_sha1_file(filename)
        line = (filename, sha1)
        liststring.append(line)
    return liststring


def save_list_csv(filename, liststring):
    """Save the list into a csv file"""
    with open(filename, "w") as targetfile:
        for item in liststring:
            filename = item[0]
            sha1 = item[1]
            line = filename + "," + sha1 + "\n"
            targetfile.write(line)


def save_list_doc(filename, liststring):
    """Save the list as a table in a word document"""
    doc = Document()
    table = doc.add_table(0, 0)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for index, item in enumerate(liststring):
        table.add_row()
        row = table.rows[index]
        filename = item[0]
        sha1 = item[1]
        print(filename + "," + sha1 + "\n")
        row.cells[0].text = str(filename)
        row.cells[1].text = str(sha1)
    doc.save(filename)


def main():
    hashes = generate_list_hash(".")
    save_list_csv("mainhashtable.csv", hashes)
    save_list_doc("mainhashtable.docx", hashes)


if __name__ == "__main__":
    main()